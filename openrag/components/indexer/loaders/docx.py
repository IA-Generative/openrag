import re
import zipfile
from io import BytesIO

from docx import Document as DocxDocument
from langchain_core.documents.base import Document
from markitdown import MarkItDown
from PIL import Image
from utils.logger import get_logger

from .base import BaseLoader, ensure_png_compatible_mode

logger = get_logger()


def convert_to_png_image(image: Image.Image) -> Image.Image:
    image = ensure_png_compatible_mode(image)
    with BytesIO() as buffer:
        image.save(buffer, format="PNG")
        buffer.seek(0)
        # Reload the image from the buffer as a PNG
        png_image = Image.open(buffer).convert("RGBA")
    return png_image


class DocxLoader(BaseLoader):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.converter = MarkItDown()

    async def aload_document(self, file_path, metadata, save_markdown=False):
        try:
            result = self.converter.convert(file_path).text_content
        except Exception as markitdown_err:
            logger.warning(
                "MarkItDown conversion failed, falling back to python-docx plain text extraction",
                path=str(file_path),
                error=str(markitdown_err),
            )
            try:
                result = self._fallback_extract_text(file_path)
            except Exception as docx_err:
                raise RuntimeError(
                    f"DOCX conversion failed with both MarkItDown ({markitdown_err}) and python-docx ({docx_err})"
                ) from docx_err

        if self.image_captioning:
            # Handle embedded images (extracted from docx zip)
            # images may contain None entries for unsupported formats (e.g. EMF, WMF)
            images = self.get_images_from_zip(file_path)
            valid_images = [img for img in images if img is not None]
            captions = await self.caption_images(valid_images, desc="Captioning embedded images")

            # Rebuild caption list preserving positional alignment with markdown refs
            caption_iter = iter(captions)
            for img in images:
                caption = next(caption_iter) if img is not None else ""
                result = re.sub(
                    r"!\[.*?\]\(data:image/.*?\)",
                    caption.replace("\\", "/") if caption else "",
                    string=result,
                    count=1,
                )

            # Handle linked images (HTTP URLs) using shared method
            # Only caption HTTP URLs, data URIs are already handled above
            result = await self.replace_markdown_images_with_captions(
                result,
                caption_data_uris=False,
                desc="Captioning linked images",
            )
        else:
            logger.info("Image captioning disabled. Ignoring images.")

        doc = Document(page_content=result, metadata=metadata)
        if save_markdown:
            self.save_content(result, str(file_path))
        return doc

    def _fallback_extract_text(self, file_path) -> str:
        doc = DocxDocument(file_path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def get_images_from_zip(self, input_file):
        try:
            docx = zipfile.ZipFile(input_file, "r")
        except zipfile.BadZipFile:
            logger.warning("File is not a valid zip archive; skipping image extraction.", path=str(input_file))
            return []
        with docx:
            file_names = docx.namelist()
            # word/media/ may also contain non-image files (e.g. oleObject, hdphoto, ink)
            image_files = [f for f in file_names if f.startswith("word/media/")]
            if not image_files:
                return []

            images_not_in_order, order = [], []

            # the images got from the original file is not in the right order
            # but the target_ref contains the position of the image in the document

            for image_file in image_files:
                image_data = docx.read(image_file)
                image_extension = image_file.split(".")[-1].lower()
                try:
                    image = Image.open(BytesIO(image_data))
                    image = convert_to_png_image(image)
                    order_num = int(image_file.split("media/image")[1].split(f".{image_extension}")[0])
                except Exception as e:
                    logger.warning(f"Skipping unsupported media file {image_file}: {e}")
                    continue

                images_not_in_order.append(image)
                order.append(order_num)

            if not images_not_in_order:
                return []

            # Reorder images by their original position in the document
            max_order = max(order)
            images = [None] * max_order
            for i, pos in enumerate(order):
                images[pos - 1] = images_not_in_order[i]
            return images
